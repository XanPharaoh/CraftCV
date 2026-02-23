import io

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


TEMPLATES = ("professional", "modern", "minimal")


def generate_resume_docx(
    *,
    full_name: str = "",
    current_title: str = "",
    location: str = "",
    education: str = "",
    skills: list[str] | None = None,
    target_role: str = "",
    bullets: list[str] | None = None,
    cover_letter: str = "",
    template: str = "professional",
    professional_summary: str = "",
    experience: list[dict] | None = None,
) -> bytes:
    """Generate a production-ready resume DOCX with no AI/tool mentions."""
    skills = skills or []
    bullets = bullets or []
    experience = experience or []

    builders = {
        "professional": _build_professional,
        "modern": _build_modern,
        "minimal": _build_minimal,
    }
    doc = builders.get(template, _build_professional)(
        full_name=full_name or "Your Name",
        title=current_title or target_role,
        location=location,
        education=education,
        skills=skills,
        target_role=target_role,
        bullets=bullets,
        cover_letter=cover_letter,
        professional_summary=professional_summary,
        experience=experience,
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Shared helpers ────────────────────────────────────────────────────────────

def _set_margins(doc, top=0.5, bottom=0.5, left=0.75, right=0.75):
    for s in doc.sections:
        s.top_margin = Inches(top)
        s.bottom_margin = Inches(bottom)
        s.left_margin = Inches(left)
        s.right_margin = Inches(right)


def _add_bottom_border(paragraph, color="CCCCCC", sz="6"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_cell_border(cell, **kwargs):
    """Set borders on a table cell. kwargs: top, bottom, left, right with dict(val, sz, color)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, attrs in kwargs.items():
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), attrs.get("val", "single"))
        el.set(qn("w:sz"), attrs.get("sz", "4"))
        el.set(qn("w:color"), attrs.get("color", "000000"))
        el.set(qn("w:space"), "0")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_cell_shading(cell, color):
    """Set background shading on a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _remove_cell_borders(cell):
    """Remove all borders from a table cell."""
    _set_cell_border(
        cell,
        top={"val": "nil", "sz": "0", "color": "FFFFFF"},
        bottom={"val": "nil", "sz": "0", "color": "FFFFFF"},
        left={"val": "nil", "sz": "0", "color": "FFFFFF"},
        right={"val": "nil", "sz": "0", "color": "FFFFFF"},
    )


def _add_run(paragraph, text, font_name="Calibri", size=10, color=None,
             bold=False, italic=False):
    """Helper to add a styled run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def _add_section_header(doc, text, font_name="Calibri", size=10.5,
                        color=RGBColor(0x1A, 0x1A, 0x1A), border_color="2C3E50",
                        uppercase=True, spacing_before=12, spacing_after=6):
    """Add a section header with underline."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(spacing_before)
    p.paragraph_format.space_after = Pt(spacing_after)
    _add_run(p, text.upper() if uppercase else text, font_name, size, color, bold=True)
    _add_bottom_border(p, color=border_color, sz="6")
    return p


def _add_bullet_item(doc, text, font_name="Calibri", size=10,
                     color=RGBColor(0x33, 0x33, 0x33)):
    """Add a properly formatted bullet point."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = Pt(14)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    _add_run(p, "•  ", font_name, size, RGBColor(0x99, 0x99, 0x99))
    _add_run(p, text, font_name, size, color)
    return p


def _add_cover_letter_page(doc, cover_letter, font_name="Calibri"):
    if not cover_letter.strip():
        return
    doc.add_page_break()

    # Cover letter header
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    _add_run(p, "Cover Letter", font_name, 14, RGBColor(0x2C, 0x3E, 0x50), bold=True)
    _add_bottom_border(p, color="2C3E50", sz="6")

    for para_text in cover_letter.split("\n"):
        stripped = para_text.strip()
        if not stripped:
            continue
        p = doc.add_paragraph()
        _add_run(p, stripped, font_name, 10.5, RGBColor(0x33, 0x33, 0x33))
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = Pt(16)


def _add_skills_grid(doc, skills, font_name="Calibri", accent=None):
    """Add skills as a clean multi-column layout."""
    if not skills:
        return
    # Create a table with invisible borders for grid layout
    cols = 3
    rows = (len(skills) + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, skill in enumerate(skills):
        row_idx = i // cols
        col_idx = i % cols
        cell = table.cell(row_idx, col_idx)
        _remove_cell_borders(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        bullet_color = accent or RGBColor(0x2C, 0x3E, 0x50)
        _add_run(p, "▸ ", font_name, 9, bullet_color)
        _add_run(p, skill, font_name, 9.5, RGBColor(0x33, 0x33, 0x33))

    # Clear borders on empty cells
    for i in range(len(skills), rows * cols):
        row_idx = i // cols
        col_idx = i % cols
        _remove_cell_borders(table.cell(row_idx, col_idx))


# ── Professional template ─────────────────────────────────────────────────────

def _build_professional(
    *, full_name, title, location, education, skills,
    target_role, bullets, cover_letter,
    professional_summary="", experience=None,
):
    doc = Document()
    _set_margins(doc, top=0.45, bottom=0.45, left=0.7, right=0.7)
    fn = "Calibri"
    dark = RGBColor(0x2C, 0x3E, 0x50)
    ink = RGBColor(0x33, 0x33, 0x33)
    gray = RGBColor(0x7F, 0x8C, 0x8D)

    # ── Name ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    _add_run(p, full_name.upper(), fn, 20, dark, bold=True)

    # ── Contact line ──
    parts = [x for x in [title, location] if x.strip()]
    if parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        for i, part in enumerate(parts):
            if i > 0:
                _add_run(p, "  ·  ", fn, 10, gray)
            _add_run(p, part, fn, 10, gray)

    # ── Divider ──
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(4)
    _add_bottom_border(p, color="2C3E50", sz="10")

    # ── Professional Summary ──
    if professional_summary:
        _add_section_header(doc, "Professional Summary", fn, border_color="2C3E50")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(15)
        _add_run(p, professional_summary, fn, 10, ink)
    elif target_role:
        _add_section_header(doc, "Professional Summary", fn, border_color="2C3E50")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        summary = f"Results-driven {title or 'professional'} with a proven track record of delivering impactful results. Seeking to leverage expertise in a {target_role} capacity."
        _add_run(p, summary, fn, 10, ink)

    # ── Core Competencies ──
    if skills:
        _add_section_header(doc, "Core Competencies", fn, border_color="2C3E50")
        _add_skills_grid(doc, skills, fn, dark)
        # Spacer
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)

    # ── Professional Experience ──
    if experience:
        _add_section_header(doc, "Professional Experience", fn, border_color="2C3E50")
        for entry in experience:
            job_title = entry.get("job_title", "")
            company = entry.get("company", "")
            dates = entry.get("dates", "")
            role_bullets = entry.get("bullets", [])

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(6)
            _add_run(p, job_title, fn, 11, dark, bold=True)

            if company or dates:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                parts = [x for x in [company, dates] if x]
                _add_run(p, "  |  ".join(parts), fn, 9.5, gray, italic=True)

            for bullet in role_bullets:
                _add_bullet_item(doc, bullet, fn, 10, ink)
    elif bullets:
        _add_section_header(doc, "Professional Experience", fn, border_color="2C3E50")

        role = target_role or title
        if role:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            _add_run(p, role, fn, 11, dark, bold=True)

        for bullet in bullets:
            _add_bullet_item(doc, bullet, fn, 10, ink)

    # ── Education ──
    if education.strip():
        _add_section_header(doc, "Education", fn, border_color="2C3E50")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _add_run(p, education, fn, 10, ink)

    _add_cover_letter_page(doc, cover_letter, fn)
    return doc


# ── Modern template ───────────────────────────────────────────────────────────

def _build_modern(
    *, full_name, title, location, education, skills,
    target_role, bullets, cover_letter,
    professional_summary="", experience=None,
):
    doc = Document()
    _set_margins(doc, top=0.4, bottom=0.4, left=0.65, right=0.65)
    fn = "Calibri"
    accent = RGBColor(0x1A, 0x6B, 0xDB)
    dark = RGBColor(0x22, 0x22, 0x22)
    ink = RGBColor(0x33, 0x33, 0x33)
    light_gray = RGBColor(0x95, 0xA5, 0xA6)

    # ── Header block with accent color bar ──
    # Use a table to create a colored header section
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = header_table.cell(0, 0)
    _set_cell_shading(cell, "1A6BDB")
    _remove_cell_borders(cell)

    # Name in header
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    _add_run(p, full_name, fn, 22, RGBColor(0xFF, 0xFF, 0xFF), bold=True)

    # Subtitle in header
    parts = [x for x in [title, location] if x.strip()]
    if parts:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        for i, part in enumerate(parts):
            if i > 0:
                _add_run(p, "  |  ", fn, 10, RGBColor(0xBD, 0xD5, 0xF7))
            _add_run(p, part, fn, 10, RGBColor(0xEC, 0xF0, 0xF1))

    # Spacer after header
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)

    # ── Professional Summary ──
    if professional_summary:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _add_run(p, "Summary", fn, 12, accent, bold=True)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = Pt(15)
        _add_run(p, professional_summary, fn, 10, ink)

    # ── Skills as tags ──
    if skills:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _add_run(p, "Skills", fn, 12, accent, bold=True)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        _add_run(p, "  •  ".join(skills), fn, 9.5, ink)

    # ── Experience ──
    if experience:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(4)
        _add_run(p, "Experience", fn, 12, accent, bold=True)
        _add_bottom_border(p, color="1A6BDB", sz="4")

        for entry in experience:
            job_title = entry.get("job_title", "")
            company = entry.get("company", "")
            dates = entry.get("dates", "")
            role_bullets = entry.get("bullets", [])

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(6)
            _add_run(p, job_title, fn, 11, dark, bold=True)

            if company or dates:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                parts = [x for x in [company, dates] if x]
                _add_run(p, "  |  ".join(parts), fn, 10, light_gray)

            for bullet in role_bullets:
                _add_bullet_item(doc, bullet, fn, 10, ink)
    elif bullets:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(4)
        _add_run(p, "Experience", fn, 12, accent, bold=True)
        _add_bottom_border(p, color="1A6BDB", sz="4")

        role = target_role or title
        if role:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.space_before = Pt(4)
            _add_run(p, role, fn, 11, dark, bold=True)

        for bullet in bullets:
            _add_bullet_item(doc, bullet, fn, 10, ink)

    # ── Education ──
    if education.strip():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        _add_run(p, "Education", fn, 12, accent, bold=True)
        _add_bottom_border(p, color="1A6BDB", sz="4")

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _add_run(p, education, fn, 10, ink)

    _add_cover_letter_page(doc, cover_letter, fn)
    return doc


# ── Minimal template ──────────────────────────────────────────────────────────

def _build_minimal(
    *, full_name, title, location, education, skills,
    target_role, bullets, cover_letter,
    professional_summary="", experience=None,
):
    doc = Document()
    _set_margins(doc, top=0.55, bottom=0.55, left=0.8, right=0.8)
    fn = "Cambria"
    black = RGBColor(0x1A, 0x1A, 0x1A)
    dark_gray = RGBColor(0x44, 0x44, 0x44)
    mid_gray = RGBColor(0x77, 0x77, 0x77)

    # ── Name — elegant, left aligned ──
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    _add_run(p, full_name, fn, 18, black, bold=True)

    # ── Contact line ──
    parts = [x for x in [title, location] if x.strip()]
    if parts:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _add_run(p, "  |  ".join(parts), fn, 9.5, mid_gray)

    # Thin rule
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    _add_bottom_border(p, color="CCCCCC", sz="2")

    # ── Professional Summary ──
    if professional_summary:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _add_run(p, "Summary", fn, 10.5, black, bold=True)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = Pt(15)
        _add_run(p, professional_summary, fn, 9.5, dark_gray)

    # ── Skills ──
    if skills:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _add_run(p, "Skills", fn, 10.5, black, bold=True)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        _add_run(p, ", ".join(skills), fn, 9.5, dark_gray)

    # ── Experience ──
    if experience:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _add_run(p, "Experience", fn, 10.5, black, bold=True)

        for entry in experience:
            job_title = entry.get("job_title", "")
            company = entry.get("company", "")
            dates = entry.get("dates", "")
            role_bullets = entry.get("bullets", [])

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(6)
            _add_run(p, job_title, fn, 10, dark_gray, italic=True)

            if company or dates:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                parts = [x for x in [company, dates] if x]
                _add_run(p, "  |  ".join(parts), fn, 9, mid_gray)

            for bullet in role_bullets:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.line_spacing = Pt(14)
                p.paragraph_format.left_indent = Inches(0.2)
                p.paragraph_format.first_line_indent = Inches(-0.15)
                _add_run(p, "—  ", fn, 9.5, mid_gray)
                _add_run(p, bullet, fn, 9.5, dark_gray)
    elif bullets:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _add_run(p, "Experience", fn, 10.5, black, bold=True)

        role = target_role or title
        if role:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            _add_run(p, role, fn, 10, dark_gray, italic=True)

        for bullet in bullets:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.line_spacing = Pt(14)
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            _add_run(p, "—  ", fn, 9.5, mid_gray)
            _add_run(p, bullet, fn, 9.5, dark_gray)

    # ── Education ──
    if education.strip():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        _add_run(p, "Education", fn, 10.5, black, bold=True)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _add_run(p, education, fn, 9.5, dark_gray)

    _add_cover_letter_page(doc, cover_letter, fn)
    return doc
